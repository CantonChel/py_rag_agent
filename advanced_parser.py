"""
高级文档解析模块

支持多种文件格式，统一转换为Markdown格式进行处理

【文件解析流程】

┌──────────────────────────────────────────────────────────────────────┐
│                        文件解析与转换流程                             │
├──────────────────────────────────────────────────────────────────────┤
│                                                                      │
│   PDF ──────┐                                                        │
│             │                                                        │
│   Word ─────┤      ┌─────────────┐      ┌─────────────┐             │
│             │  →   │  文件解析器  │  →   │  Markdown   │  →  清洗    │
│   HTML ─────┤      │             │      │  转换器     │             │
│             │      └─────────────┘      └─────────────┘             │
│   PPT ──────┤                                                        │
│             │                                                        │
│   MD/TXT ───┘                                                        │
│                                                                      │
│   【各格式处理要点】                                                 │
│                                                                      │
│   PDF:                                                               │
│   ├── 文本提取: pdfplumber / PyMuPDF                               │
│   ├── 表格识别: camelot / tabula                                   │
│   ├── OCR处理: pytesseract / PaddleOCR (扫描件)                    │
│   ├── 图片提取: 提取嵌入图片                                        │
│   └── 布局分析: 保持段落结构                                        │
│                                                                      │
│   Word (docx):                                                       │
│   ├── 样式转换: 标题→##, 加粗→**等                                 │
│   ├── 表格处理: 转为Markdown表格                                    │
│   ├── 图片提取: 提取并保存                                          │
│   └── 目录结构: 保持层级                                            │
│                                                                      │
│   HTML:                                                              │
│   ├── 标签清理: 去除script/style等                                 │
│   ├── 结构转换: h1→#, p保持, 列表转换                              │
│   ├── 链接处理: 保留或转为引用                                      │
│   └── 代码块: 识别并保留                                            │
│                                                                      │
│   PPT:                                                               │
│   ├── 幻灯片遍历: 逐页提取                                          │
│   ├── 文本提取: 按文本框组织                                        │
│   ├── 图片提取: 提取幻灯片图片                                      │
│   └── 备注: 提取演讲者备注                                          │
│                                                                      │
└──────────────────────────────────────────────────────────────────────┘
"""

import os
import re
import base64
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, field
from abc import ABC, abstractmethod


@dataclass
class ImageInfo:
    """
    图片信息数据结构
    
    存储从文档中提取的图片的相关信息
    """
    image_id: str
    source_path: str
    page_number: Optional[int] = None
    image_data: Optional[bytes] = None
    image_path: Optional[str] = None
    caption: Optional[str] = None
    surrounding_text: Optional[str] = None
    related_chunks: List[str] = field(default_factory=list)
    
    def to_base64(self) -> str:
        """将图片数据转换为base64字符串"""
        if self.image_data:
            return base64.b64encode(self.image_data).decode('utf-8')
        return ""
    
    def get_hash(self) -> str:
        """获取图片内容的hash值（用于去重）"""
        if self.image_data:
            return hashlib.md5(self.image_data).hexdigest()
        return self.image_id


@dataclass
class ParsedDocument:
    """
    解析后的文档结构
    
    【数据结构说明】
    这是文档解析后的统一输出格式，包含：
    - 文本内容（Markdown格式）
    - 元数据信息
    - 提取的图片
    - 结构信息（章节、表格等）
    """
    doc_id: str
    file_path: str
    file_type: str
    markdown_content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    images: List[ImageInfo] = field(default_factory=list)
    sections: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    
    # 关系映射
    chunk_to_image: Dict[str, List[str]] = field(default_factory=dict)
    image_to_chunk: Dict[str, List[str]] = field(default_factory=dict)


class BaseParser(ABC):
    """
    文件解析器基类
    
    所有具体格式的解析器都需要继承此类并实现parse方法
    """
    
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文件并返回统一结构的文档
        
        参数:
            file_path: 文件路径
            
        返回:
            解析后的文档对象
        """
        pass
    
    @staticmethod
    def generate_doc_id(file_path: str) -> str:
        """生成文档唯一ID"""
        return hashlib.md5(file_path.encode()).hexdigest()[:12]


class MarkdownCleaner:
    """
    Markdown文本清洗器
    
    【清洗步骤详解】
    
    步骤1: 去除控制字符
    ├── 移除不可见的ASCII控制字符
    └── 保留换行符和制表符
    
    步骤2: 标准化空白
    ├── 多个空格→单个空格
    ├── 多个空行→两个空行
    └── 去除行首行尾空白
    
    步骤3: 修复Markdown格式
    ├── 确保标题前后有空行
    ├── 修复列表格式
    └── 处理代码块
    
    步骤4: 去除噪音
    ├── 去除页码
    ├── 去除页眉页脚
    └── 去除水印文字
    """
    
    @staticmethod
    def clean(markdown_text: str) -> str:
        """
        执行完整的清洗流程
        
        参数:
            markdown_text: 原始Markdown文本
            
        返回:
            清洗后的文本
        """
        if not markdown_text:
            return ""
        
        text = markdown_text
        
        # 步骤1: 去除控制字符
        text = MarkdownCleaner._remove_control_chars(text)
        
        # 步骤2: 标准化空白
        text = MarkdownCleaner._normalize_whitespace(text)
        
        # 步骤3: 修复Markdown格式
        text = MarkdownCleaner._fix_markdown_format(text)
        
        # 步骤4: 去除噪音
        text = MarkdownCleaner._remove_noise(text)
        
        return text.strip()
    
    @staticmethod
    def _remove_control_chars(text: str) -> str:
        """去除控制字符"""
        # 保留换行(\n)、制表符(\t)、回车(\r)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text
    
    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        """标准化空白字符"""
        # 多个空格→单个空格（保留Markdown缩进）
        lines = text.split('\n')
        normalized_lines = []
        
        for line in lines:
            # 如果是代码块或列表，保留原始缩进
            if line.startswith('    ') or line.startswith('\t') or line.startswith(('- ', '* ', '1. ', '2. ')):
                normalized_lines.append(line)
            else:
                # 普通行：多个空格→单个空格
                normalized_lines.append(re.sub(r'[ \t]+', ' ', line))
        
        text = '\n'.join(normalized_lines)
        
        # 多个空行→两个空行
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        
        return text
    
    @staticmethod
    def _fix_markdown_format(text: str) -> str:
        """修复Markdown格式"""
        lines = text.split('\n')
        fixed_lines = []
        
        for i, line in enumerate(lines):
            # 确保标题前后有空行
            if line.startswith('#'):
                if i > 0 and fixed_lines and fixed_lines[-1].strip():
                    fixed_lines.append('')
                fixed_lines.append(line)
                if i < len(lines) - 1 and lines[i + 1].strip() and not lines[i + 1].startswith('#'):
                    fixed_lines.append('')
            else:
                fixed_lines.append(line)
        
        return '\n'.join(fixed_lines)
    
    @staticmethod
    def _remove_noise(text: str) -> str:
        """去除噪音内容"""
        # 去除页码
        text = re.sub(r'^\s*[-–—]?\s*\d+\s*[-–—]?\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^第\s*\d+\s*页', '', text, flags=re.MULTILINE)
        
        # 去除常见的水印文字（可根据需要添加）
        watermark_patterns = [
            r'仅限.*?使用',
            r'内部资料.*?禁止外传',
            r'Confidential',
        ]
        for pattern in watermark_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        return text


class PDFParser(BaseParser):
    """
    PDF文件解析器
    
    【PDF解析流程】
    
    1. 检测PDF类型
       ├── 文本型PDF: 直接提取文本
       └── 扫描型PDF: 需要OCR处理
       
    2. 提取内容
       ├── 文本: 按页面/段落提取
       ├── 表格: 识别表格结构
       ├── 图片: 提取嵌入图片
       └── 布局: 保持原有结构
       
    3. 转换为Markdown
       ├── 标题识别: 根据字体大小判断
       ├── 段落合并: 保持段落完整
       ├── 表格转换: 转为MD表格格式
       └── 图片引用: 插入图片标记
    
    【依赖库】
    - pdfplumber: 文本和表格提取
    - PyMuPDF (fitz): 图片提取
    - pytesseract/PaddleOCR: OCR（扫描件）
    """
    
    def __init__(self, extract_images: bool = True, use_ocr: bool = False):
        """
        初始化PDF解析器
        
        参数:
            extract_images: 是否提取图片
            use_ocr: 是否使用OCR处理扫描件
        """
        self.extract_images = extract_images
        self.use_ocr = use_ocr
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析PDF文件
        
        【处理步骤】
        1. 打开PDF文件
        2. 逐页提取文本、表格、图片
        3. 识别标题层级
        4. 组装Markdown内容
        5. 收集图片信息
        6. 构建元数据
        
        参数:
            file_path: PDF文件路径
            
        返回:
            解析后的文档对象
        """
        doc_id = self.generate_doc_id(file_path)
        
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("请安装pdfplumber: pip install pdfplumber")
        
        markdown_parts = []
        images = []
        tables = []
        sections = []
        
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                # 添加页面分隔
                markdown_parts.append(f"\n\n---\n\n**第 {page_num} 页**\n\n")
                
                # 提取文本
                text = page.extract_text() or ""
                
                # 尝试识别标题（基于字体大小）
                text = self._detect_headers(text, page)
                
                markdown_parts.append(text)
                
                # 提取表格
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    md_table = self._table_to_markdown(table)
                    markdown_parts.append(f"\n\n{md_table}\n")
                    tables.append({
                        "page": page_num,
                        "index": table_idx,
                        "content": md_table
                    })
                
                # 提取图片
                if self.extract_images:
                    page_images = self._extract_images_from_page(page, page_num)
                    for img in page_images:
                        images.append(img)
                        markdown_parts.append(f"\n\n![图片](image://{img.image_id})\n")
        
        markdown_content = '\n'.join(markdown_parts)
        markdown_content = MarkdownCleaner.clean(markdown_content)
        
        return ParsedDocument(
            doc_id=doc_id,
            file_path=file_path,
            file_type="pdf",
            markdown_content=markdown_content,
            metadata={
                "total_pages": total_pages,
                "has_images": len(images) > 0,
                "has_tables": len(tables) > 0
            },
            images=images,
            sections=sections,
            tables=tables
        )
    
    def _detect_headers(self, text: str, page) -> str:
        """
        检测并标记标题
        
        【标题识别策略】
        1. 基于字体大小：大字体→标题
        2. 基于格式：粗体、居中→可能是标题
        3. 基于内容：数字开头（如"1. xxx"）→可能是标题
        """
        # 简化实现：通过正则识别常见的标题模式
        lines = text.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测数字编号的标题 (如 "1. xxx", "1.1 xxx")
            if re.match(r'^\d+\.\d*\s+.+', line):
                # 判断是一级还是二级标题
                if re.match(r'^\d+\.\s+.+', line):
                    line = f"# {line}"
                else:
                    line = f"## {line}"
            
            # 检测中文数字标题 (如 "一、xxx")
            elif re.match(r'^[一二三四五六七八九十]+[、.]\s*.+', line):
                line = f"# {line}"
            
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)
    
    def _table_to_markdown(self, table: List[List[str]]) -> str:
        """
        将表格数据转换为Markdown格式
        
        参数:
            table: 二维表格数据
            
        返回:
            Markdown表格字符串
        """
        if not table:
            return ""
        
        # 清理表格单元格
        cleaned_table = []
        for row in table:
            cleaned_row = [cell.replace('\n', ' ') if cell else "" for cell in row]
            cleaned_table.append(cleaned_row)
        
        # 生成Markdown表格
        md_lines = []
        
        # 表头
        header = cleaned_table[0]
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        
        # 数据行
        for row in cleaned_table[1:]:
            md_lines.append("| " + " | ".join(row) + " |")
        
        return '\n'.join(md_lines)
    
    def _extract_images_from_page(self, page, page_num: int) -> List[ImageInfo]:
        """
        从PDF页面提取图片
        
        参数:
            page: pdfplumber页面对象
            page_num: 页码
            
        返回:
            图片信息列表
        """
        images = []
        
        # 获取页面中的图片对象
        # 注意：pdfplumber对图片提取支持有限
        # 实际项目中建议使用PyMuPDF (fitz)
        
        try:
            import fitz
            doc = fitz.open(page.pdf.stream.name if hasattr(page.pdf, 'stream') else "")
            fitz_page = doc[page_num - 1]
            
            image_list = fitz_page.get_images()
            
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                
                image_id = f"img_{page_num}_{img_index}"
                images.append(ImageInfo(
                    image_id=image_id,
                    source_path="",
                    page_number=page_num,
                    image_data=base_image["image"],
                    caption=None
                ))
        except Exception:
            pass
        
        return images


class WordParser(BaseParser):
    """
    Word文档解析器
    
    【Word解析流程】
    
    1. 读取文档结构
       ├── 段落: 遍历所有段落
       ├── 表格: 提取表格内容
       ├── 图片: 提取嵌入图片
       └── 样式: 获取格式信息
       
    2. 样式映射
       ├── 标题样式 → Markdown标题(#)
       ├── 加粗 → **text**
       ├── 斜体 → *text*
       ├── 列表 → - text 或 1. text
       └── 代码 → `text` 或 ```code```
       
    3. 图片处理
       ├── 提取图片数据
       ├── 保存到指定目录
       └── 生成引用标记
    """
    
    def __init__(self, image_output_dir: str = "./extracted_images"):
        """
        初始化Word解析器
        
        参数:
            image_output_dir: 图片输出目录
        """
        self.image_output_dir = image_output_dir
        Path(image_output_dir).mkdir(parents=True, exist_ok=True)
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析Word文档
        
        参数:
            file_path: Word文件路径(.docx)
            
        返回:
            解析后的文档对象
        """
        doc_id = self.generate_doc_id(file_path)
        
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装python-docx: pip install python-docx")
        
        doc = Document(file_path)
        markdown_parts = []
        images = []
        tables = []
        sections = []
        
        # 处理段落
        for para in doc.paragraphs:
            md_text = self._paragraph_to_markdown(para)
            if md_text:
                markdown_parts.append(md_text)
                
                # 检测标题，记录章节
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.replace('Heading ', '').replace('标题 ', '')) if para.style.name[-1].isdigit() else 1
                    sections.append({
                        "level": level,
                        "title": para.text
                    })
        
        # 处理表格
        for table_idx, table in enumerate(doc.tables):
            md_table = self._table_to_markdown(table)
            markdown_parts.append(f"\n\n{md_table}\n")
            tables.append({
                "index": table_idx,
                "content": md_table
            })
        
        # 提取图片
        images = self._extract_images(doc, doc_id)
        
        markdown_content = '\n'.join(markdown_parts)
        markdown_content = MarkdownCleaner.clean(markdown_content)
        
        return ParsedDocument(
            doc_id=doc_id,
            file_path=file_path,
            file_type="docx",
            markdown_content=markdown_content,
            metadata={
                "has_images": len(images) > 0,
                "has_tables": len(tables) > 0
            },
            images=images,
            sections=sections,
            tables=tables
        )
    
    def _paragraph_to_markdown(self, para) -> str:
        """
        将Word段落转换为Markdown
        
        【转换规则】
        - Heading 1 → #
        - Heading 2 → ##
        - List Bullet → -
        - List Number → 1. 2. 3.
        - 加粗 → **text**
        - 斜体 → *text*
        """
        text = para.text.strip()
        if not text:
            return ""
        
        style_name = para.style.name
        
        # 处理标题
        if 'Heading' in style_name or '标题' in style_name:
            level = 1
            if style_name[-1].isdigit():
                level = int(style_name[-1])
            return f"{'#' * level} {text}"
        
        # 处理列表
        if 'List Bullet' in style_name:
            return f"- {text}"
        if 'List Number' in style_name:
            return f"1. {text}"
        
        # 处理普通段落（检查内部样式）
        runs_text = []
        for run in para.runs:
            run_text = run.text
            if run.bold and run.italic:
                run_text = f"***{run_text}***"
            elif run.bold:
                run_text = f"**{run_text}**"
            elif run.italic:
                run_text = f"*{run_text}*"
            runs_text.append(run_text)
        
        return ''.join(runs_text) if runs_text else text
    
    def _table_to_markdown(self, table) -> str:
        """将Word表格转换为Markdown"""
        rows = []
        for row in table.rows:
            cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            rows.append(cells)
        
        if not rows:
            return ""
        
        md_lines = []
        md_lines.append("| " + " | ".join(rows[0]) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(row) + " |")
        
        return '\n'.join(md_lines)
    
    def _extract_images(self, doc, doc_id: str) -> List[ImageInfo]:
        """从Word文档提取图片"""
        images = []
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_id = f"{doc_id}_img_{len(images)}"
                images.append(ImageInfo(
                    image_id=image_id,
                    source_path="",
                    image_data=image_data
                ))
        
        return images


class HTMLParser(BaseParser):
    """
    HTML文件解析器
    
    【HTML解析流程】
    
    1. 解析HTML结构
       ├── BeautifulSoup解析DOM
       └── 提取body内容
       
    2. 清理无关内容
       ├── 移除script/style标签
       ├── 移除注释
       └── 移除导航/页脚等
       
    3. 转换为Markdown
       ├── h1-h6 → #-######
       ├── p → 段落
       ├── ul/ol → 列表
       ├── table → MD表格
       ├── code/pre → 代码块
       └── a → 链接
       
    4. 处理图片
       ├── 下载远程图片
       └── 提取base64图片
    """
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析HTML文件
        
        参数:
            file_path: HTML文件路径
            
        返回:
            解析后的文档对象
        """
        doc_id = self.generate_doc_id(file_path)
        
        try:
            from bs4 import BeautifulSoup
            import html2text
        except ImportError:
            raise ImportError("请安装依赖: pip install beautifulsoup4 html2text")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 移除不需要的标签
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            tag.decompose()
        
        # 移除注释
        from bs4 import Comment
        for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
            comment.extract()
        
        # 使用html2text转换
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.body_width = 0
        
        markdown_content = h.handle(str(soup))
        markdown_content = MarkdownCleaner.clean(markdown_content)
        
        # 提取图片
        images = []
        for idx, img in enumerate(soup.find_all('img')):
            image_id = f"{doc_id}_img_{idx}"
            images.append(ImageInfo(
                image_id=image_id,
                source_path=img.get('src', ''),
                caption=img.get('alt', '')
            ))
        
        return ParsedDocument(
            doc_id=doc_id,
            file_path=file_path,
            file_type="html",
            markdown_content=markdown_content,
            metadata={
                "has_images": len(images) > 0
            },
            images=images
        )


class MarkdownParser(BaseParser):
    """
    Markdown文件解析器
    
    Markdown文件相对简单，主要是读取和清洗
    """
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析Markdown文件
        
        参数:
            file_path: Markdown文件路径
            
        返回:
            解析后的文档对象
        """
        doc_id = self.generate_doc_id(file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 清洗内容
        cleaned_content = MarkdownCleaner.clean(content)
        
        # 提取章节结构
        sections = []
        for line in cleaned_content.split('\n'):
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                sections.append({
                    "level": level,
                    "title": title
                })
        
        # 提取图片引用
        images = []
        img_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        for idx, match in enumerate(re.finditer(img_pattern, cleaned_content)):
            image_id = f"{doc_id}_img_{idx}"
            images.append(ImageInfo(
                image_id=image_id,
                source_path=match.group(2),
                caption=match.group(1)
            ))
        
        return ParsedDocument(
            doc_id=doc_id,
            file_path=file_path,
            file_type="md",
            markdown_content=cleaned_content,
            metadata={},
            images=images,
            sections=sections
        )


class DocumentParserFactory:
    """
    文档解析器工厂
    
    【工厂模式】
    根据文件类型自动选择对应的解析器
    """
    
    _parsers: Dict[str, type] = {
        '.pdf': PDFParser,
        '.docx': WordParser,
        '.doc': WordParser,
        '.html': HTMLParser,
        '.htm': HTMLParser,
        '.md': MarkdownParser,
        '.markdown': MarkdownParser,
        '.txt': MarkdownParser,
    }
    
    @classmethod
    def get_parser(cls, file_path: str, **kwargs) -> BaseParser:
        """
        获取对应的解析器
        
        参数:
            file_path: 文件路径
            **kwargs: 解析器参数
            
        返回:
            解析器实例
            
        异常:
            ValueError: 不支持的文件类型
        """
        ext = Path(file_path).suffix.lower()
        
        if ext not in cls._parsers:
            raise ValueError(f"不支持的文件类型: {ext}")
        
        parser_class = cls._parsers[ext]
        return parser_class(**kwargs)
    
    @classmethod
    def parse(cls, file_path: str, **kwargs) -> ParsedDocument:
        """
        解析文件（便捷方法）
        
        参数:
            file_path: 文件路径
            **kwargs: 解析器参数
            
        返回:
            解析后的文档对象
        """
        parser = cls.get_parser(file_path, **kwargs)
        return parser.parse(file_path)
    
    @classmethod
    def supported_formats(cls) -> List[str]:
        """返回支持的文件格式列表"""
        return list(cls._parsers.keys())


def demo_parser():
    """演示文档解析功能"""
    print("支持的文件格式:", DocumentParserFactory.supported_formats())
    
    # 示例：解析一个markdown文件
    sample_md = """
# 示例文档
    
## 第一章
    
这是一段示例文本。
    
### 1.1 子章节
    
- 列表项1
- 列表项2
    
![示例图片](image.png)
    
| 列1 | 列2 |
|-----|-----|
| A   | B   |
"""
    
    # 创建临时文件
    import tempfile
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(sample_md)
        temp_path = f.name
    
    # 解析文件
    doc = DocumentParserFactory.parse(temp_path)
    
    print(f"\n文档ID: {doc.doc_id}")
    print(f"文件类型: {doc.file_type}")
    print(f"章节数量: {len(doc.sections)}")
    print(f"图片数量: {len(doc.images)}")
    print(f"\n内容预览:\n{doc.markdown_content[:200]}...")
    
    # 清理
    os.unlink(temp_path)


if __name__ == "__main__":
    demo_parser()
